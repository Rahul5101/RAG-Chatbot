
import fitz 
import io
from PIL import Image
# import google.generativeai as genai
from google import genai    
import os
import time
import csv
from docx import Document 
from concurrent.futures import ThreadPoolExecutor, as_completed
from dotenv import load_dotenv
load_dotenv()


# --------------------------
# 1. Configure Gemini
# --------------------------
# genai.configure(api_key=API_KEY)
# # We will use the model for OCR (PDFs only)
# model = genai.GenerativeModel("gemini-2.0-flash") 
API_KEY = os.getenv("GOOGLE_API_KEY")
client = genai.Client(api_key=API_KEY)

# --------------------------
# 2. PDF → Images (Existing Code)
# --------------------------
def pdf_to_images(pdf_path):        
    pdf_document = fitz.open(pdf_path)
    images = []
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        pix = page.get_pixmap()
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        images.append(img)
    pdf_document.close()
    return images

# --------------------------
# 3. Text Extraction for CSV and DOCX
# --------------------------

def extract_text_from_csv(file_path):
    """Extracts text content from a CSV file, formatting each row."""
    extracted_text = []
 
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader, 1):
                # Joins columns into a single line, separated by tabs or pipes
                row_text = "|".join(row)
                extracted_text.append(f"\n--- Row {i} ---\n{row_text}")
        return "\n".join(extracted_text)
    except Exception as e:
        return f"[ERROR] Failed to extract CSV: {str(e)}"

def extract_text_from_docx(file_path):
    """Extracts text content from a DOCX file."""
    doc = Document(file_path)   
    extracted_text = []
    
    # Extract paragraphs
    for i, para in enumerate(doc.paragraphs, 1):
        extracted_text.append(f"\n--- Paragraph {i} ---\n{para.text}")

    # Extract table content (optional but good practice)
    if doc.tables:
        extracted_text.append("\n--- Tables ---\n")
        for j, table in enumerate(doc.tables, 1):
            extracted_text.append(f"\n--- Table {j} ---\n")
            for row in table.rows:
                row_cells = [cell.text for cell in row.cells]
                extracted_text.append("|".join(row_cells))
                
    return "\n".join(extracted_text)

# --------------------------
# 4. OCR with Gemini (per chunk - used for PDF only)
# --------------------------
def extract_text_with_gemini(
    images,
    output_base_path,
    base_name,
    chunk_size=100  # 🔹 100 pages per file
):
    total_pages = len(images)
    results = []

    os.makedirs(output_base_path, exist_ok=True)

    for chunk_start in range(0, total_pages, chunk_size):
        chunk_end = min(chunk_start + chunk_size, total_pages)

        # 🔹 File name per chunk (page range)
        output_txt_path = os.path.join(
            output_base_path,
            f"{base_name}_pages_{chunk_start + 1}_{chunk_end}.txt"
        )

        with open(output_txt_path, "w", encoding="utf-8") as f:
            for page_index in range(chunk_start, chunk_end):
                img = images[page_index]
                page_number = page_index + 1

                time.sleep(2)  # API safety delay

                prompt = '''
                    Extract each and every content in this image as it is even if some content is repeated.
                    Explain in an understandable way and do explain the mapping if it is needed
                    (which value belongs to which field).
                    if there is handwritten words then extract them in english word precisely.
                    Provide everything in English, if the text
                    is Hindi or some other language then first translate into English.
                    Do not make any assumption in prediction, always try to extract that is present in the image.

                    You are a text extractor from images. Follow these steps carefully:

                    1. First analyze the image and then extract **all visible content exactly as it appears**, even if some content is repeated.  
                    2. If the image contains a **table**, extract it properly and also explain the mapping between rows and columns so that the semantic meaning is not lost.  
                    3. If the image contains **normal text or paragraphs**, extract them exactly as they are without adding or modifying anything.  
                    4. Only explain the mapping for tables.  
                    5. If there are handwritten words, transcribe them accurately into English.   
                    7. Do not add any extra text, introductions, or conclusions before or after the extraction. Only provide the extracted content (and table mapping if applicable).  
                    '''

                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=[prompt, img]
                )

                page_text = response.text or ""
                f.write(f"\n--- Page {page_number} ---\n{page_text}\n")
                f.flush()

                print(f"[INFO] Page {page_number} saved → {output_txt_path}")

        results.append(
            f"[CHUNK DONE] Pages {chunk_start + 1}-{chunk_end} → {output_txt_path}"
        )

    return results


# --------------------------
# 5. Process a Single File (Unified)
# --------------------------
def process_file(file_path, input_folder, output_folder):
    extension = os.path.splitext(file_path)[1].lower()
    
    # Check supported extensions
    if extension not in ['.pdf', '.csv', '.docx']:
        return f"Skipping {file_path} (Unsupported file type: {extension})."

    rel_path = os.path.relpath(file_path, input_folder)
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_subfolder = os.path.join(output_folder, os.path.dirname(rel_path))
    os.makedirs(output_subfolder, exist_ok=True)
    
    t0 = time.time()
    
    try:
        if extension == '.pdf':
            # --- PDF Processing (using Gemini/Vision) ---
            print(f"[PROCESS] Converting {file_path} to images...")
            images = pdf_to_images(file_path)
            
            # This calls the Gemini function and saves the files in chunks
            results = extract_text_with_gemini(images, output_subfolder, base_name, chunk_size=100)
            result_message = "; ".join(results)
            
        else:
            # --- Text Processing (CSV/DOCX) ---
            output_txt_path = os.path.join(output_subfolder, f"{base_name}_{extension}.txt")

            if extension == '.csv':
                print(f"[PROCESS] Extracting text from CSV: {file_path}")
                extracted_text = extract_text_from_csv(file_path)
            elif extension == '.docx':
                print(f"[PROCESS] Extracting text from DOCX: {file_path}")
                extracted_text = extract_text_from_docx(file_path)
            else:
                return f":x: Error: Unhandled extension logic for {extension}"

            # Save all extracted text to a single file
            with open(output_txt_path, "w", encoding="utf-8") as f:
                f.write(extracted_text)
                
            result_message = f"Saved text to {output_txt_path}"
            
        elapsed = time.time() - t0
        return f"✅ {file_path} processed in {elapsed:.2f} sec → {result_message}"
        
    except Exception as e:
        return f":x: Error processing {file_path}: {str(e)}"

# --------------------------
# 6. MAIN (Batch + Concurrency)
# --------------------------
def main():
    input_folder = r"D:\valiance_sol\navy-wss\demo_data"
    output_folder = r"D:\valiance_sol\navy-wss\cleaned_file"
    os.makedirs(output_folder, exist_ok=True)

    # Recursively collect all supported files
    supported_files = []
    SUPPORTED_EXTENSIONS = ['.pdf', '.csv', '.docx']
    
    for root, dirs, files in os.walk(input_folder):
        for file in files:
            extension = os.path.splitext(file)[1].lower()
            if extension in SUPPORTED_EXTENSIONS:
                supported_files.append(os.path.join(root, file))

    if not supported_files:
        print(f"⚠️ No supported files found in {input_folder} with extensions: {SUPPORTED_EXTENSIONS}")
        return

    start = time.time()
    # Run with 1 worker for safety (adjust max_workers based on your Gemini API limits and local machine specs)
    with ThreadPoolExecutor(max_workers=1) as executor:
        futures = [executor.submit(process_file, doc_file, input_folder, output_folder) for doc_file in supported_files]
        for future in as_completed(futures):
            print(future.result())
    print(f"\n:rocket: All done in {time.time() - start:.2f} sec")

if __name__ == "__main__":
    main()